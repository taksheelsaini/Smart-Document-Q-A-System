# Copyright © Taksheel Saini. All rights reserved. | GitHub: https://github.com/taksheelsaini | LinkedIn: https://www.linkedin.com/in/taksheelsaini/"""
Document processing: text extraction from PDF/DOCX and recursive text chunking.

Design rationale for chunking:
  - Recursive separator strategy splits on natural semantic boundaries first
    (paragraph breaks → sentence ends → words → characters), ensuring chunks
    land on meaningful units rather than mid-sentence cuts.
  - Overlap carries the tail of each chunk into the next, preserving context
    that would otherwise be split across a retrieval boundary.
  - All chunks are stripped and deduplicated before storage.
"""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------


def extract_text_from_pdf(file_path: Path) -> str:
    """
    Extract text from a PDF using PyMuPDF (fitz).

    PyMuPDF is preferred over pypdf because it faithfully reconstructs reading
    order, handles multi-column layouts better, and is ~10x faster on large files.
    Each page is prefixed with a page-number marker so downstream chunking can
    include page provenance in the text.
    """
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(str(file_path))
        pages: list[str] = []
        for page_num, page in enumerate(doc):
            text = page.get_text("text").strip()
            if text:
                pages.append(f"[Page {page_num + 1}]\n{text}")
        doc.close()

        if not pages:
            raise ValueError("PDF contains no extractable text (may be image-only).")

        return "\n\n".join(pages)

    except Exception as exc:
        logger.error("PDF extraction failed for %s: %s", file_path, exc)
        raise ValueError(f"Could not extract text from PDF: {exc}") from exc


def extract_text_from_docx(file_path: Path) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Heading styles are preserved with Markdown-style markers so the chunker
    can treat them as paragraph separators. Table cells are joined with pipes
    to maintain tabular structure in plain text.
    """
    try:
        from docx import Document

        doc = Document(str(file_path))
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith("Heading"):
                level = para.style.name.split()[-1] if para.style.name != "Heading" else "1"
                parts.append(f"\n{'#' * int(level)} {text}\n")
            else:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        if not parts:
            raise ValueError("DOCX appears to be empty.")

        return "\n\n".join(parts)

    except Exception as exc:
        logger.error("DOCX extraction failed for %s: %s", file_path, exc)
        raise ValueError(f"Could not extract text from DOCX: {exc}") from exc


def extract_text(file_path: Path, file_type: str) -> str:
    """Dispatch extraction to the appropriate handler based on file type."""
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type == "docx":
        return extract_text_from_docx(file_path)
    raise ValueError(f"Unsupported file type: '{file_type}'")


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

_SEPARATORS = ["\n\n", "\n", ". ", "! ", "? ", "; ", ", ", " "]


def chunk_text(text: str, chunk_size: int = 800, overlap: int = 150) -> list[str]:
    """
    Split *text* into overlapping chunks of at most *chunk_size* characters.

    Algorithm (recursive separator descent):
      1. Try splitting on the highest-priority separator that appears in text.
      2. Greedily accumulate parts until the next part would exceed chunk_size.
      3. On overflow, flush the current chunk and start the next one with
         an overlap window from the tail of the previous chunk.
      4. If any resulting chunk is still too large, recurse with the next
         separator down the priority list.
      5. Fall back to hard character-level splitting if no separator works.

    This produces chunks that almost always end on word/sentence boundaries,
    which significantly improves embedding quality and retrieval relevance.
    """

    def _split(text: str, seps: list[str]) -> list[str]:
        if len(text) <= chunk_size:
            return [text] if text.strip() else []

        sep = next((s for s in seps if s in text), "")
        remaining_seps = seps[seps.index(sep) + 1 :] if sep else []

        if not sep:
            # Hard character split as last resort
            result = []
            for i in range(0, len(text), chunk_size - overlap):
                piece = text[i : i + chunk_size].strip()
                if piece:
                    result.append(piece)
            return result

        parts = text.split(sep)
        chunks: list[str] = []
        current_parts: list[str] = []
        current_len = 0

        for part in parts:
            part_len = len(part) + len(sep)
            if current_len + part_len > chunk_size and current_parts:
                # Flush current chunk
                chunk = sep.join(current_parts).strip()
                if chunk:
                    chunks.append(chunk)

                # Build overlap prefix from tail of current_parts
                overlap_parts: list[str] = []
                overlap_len = 0
                for prev in reversed(current_parts):
                    candidate_len = len(prev) + len(sep)
                    if overlap_len + candidate_len <= overlap:
                        overlap_parts.insert(0, prev)
                        overlap_len += candidate_len
                    else:
                        break

                current_parts = overlap_parts + [part]
                current_len = sum(len(p) + len(sep) for p in current_parts)
            else:
                current_parts.append(part)
                current_len += part_len

        if current_parts:
            chunk = sep.join(current_parts).strip()
            if chunk:
                chunks.append(chunk)

        # Recurse on any chunks still exceeding the limit
        result: list[str] = []
        for ch in chunks:
            if len(ch) > chunk_size and remaining_seps:
                result.extend(_split(ch, remaining_seps))
            elif ch.strip():
                result.append(ch.strip())
        return result

    raw = _split(text, _SEPARATORS)

    # Deduplicate while preserving order
    seen: set[str] = set()
    unique: list[str] = []
    for ch in raw:
        if ch not in seen:
            seen.add(ch)
            unique.append(ch)
    return unique
