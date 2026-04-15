#!/usr/bin/env python3
"""
Generate three sample documents for testing DocQA.

Run this from the project root:
    python scripts/generate_sample_docs.py

Requires: reportlab, python-docx
    pip install reportlab python-docx
"""
import sys
from pathlib import Path

OUTPUT_DIR = Path(__file__).parent.parent / "sample_docs"
OUTPUT_DIR.mkdir(exist_ok=True)


# ── Sample 1: AI Overview PDF ─────────────────────────────────────────────────

def create_ai_overview_pdf():
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except ImportError:
        print("reportlab not installed. Run: pip install reportlab")
        return

    path = OUTPUT_DIR / "ai_overview.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=72)

    styles = getSampleStyleSheet()
    h1 = styles["Heading1"]
    h2 = styles["Heading2"]
    body = styles["BodyText"]
    body.spaceAfter = 8

    content = [
        Paragraph("Artificial Intelligence: A Comprehensive Overview", h1),
        Spacer(1, 0.2 * inch),

        Paragraph("1. What is Artificial Intelligence?", h2),
        Paragraph(
            "Artificial Intelligence (AI) refers to the simulation of human intelligence in machines "
            "programmed to think, learn, and problem-solve. The term was coined by John McCarthy in "
            "1956 at the Dartmouth Conference, widely considered the founding moment of AI as a field. "
            "AI systems process large amounts of data, identify patterns, and make decisions with "
            "minimal human intervention.", body),

        Paragraph("2. Key Branches of AI", h2),
        Paragraph(
            "Machine Learning (ML) is a subset of AI that enables systems to automatically learn and "
            "improve from experience without being explicitly programmed. ML algorithms build a model "
            "from training data and use it to make predictions or decisions.", body),
        Paragraph(
            "Deep Learning is a further subset of ML that uses artificial neural networks with many "
            "layers (hence 'deep') to model complex patterns. It powers modern image recognition, "
            "natural language processing, and speech synthesis systems.", body),
        Paragraph(
            "Natural Language Processing (NLP) enables computers to understand, interpret, and "
            "generate human language. Applications include machine translation, sentiment analysis, "
            "chatbots, and document summarization.", body),
        Paragraph(
            "Computer Vision allows machines to interpret and make decisions based on visual data "
            "from images and videos. It is used in autonomous vehicles, medical imaging, and "
            "quality control in manufacturing.", body),

        Paragraph("3. Applications of AI", h2),
        Paragraph(
            "Healthcare: AI is revolutionizing diagnostics, drug discovery, and personalized medicine. "
            "Algorithms trained on medical images can detect cancers earlier than human radiologists "
            "in some studies. IBM Watson Oncology assists oncologists in treatment planning.", body),
        Paragraph(
            "Finance: Algorithmic trading, fraud detection, credit scoring, and robo-advisors all "
            "rely on AI. JPMorgan's COiN platform can review 12,000 commercial credit agreements in "
            "seconds — a task that previously took lawyers 360,000 hours annually.", body),
        Paragraph(
            "Transportation: Self-driving vehicles from companies like Waymo and Tesla use a "
            "combination of computer vision, lidar, and reinforcement learning. Logistics companies "
            "use AI for route optimization and demand forecasting.", body),

        Paragraph("4. Ethical Considerations", h2),
        Paragraph(
            "Bias and Fairness: AI systems trained on biased datasets replicate and amplify those "
            "biases. The COMPAS recidivism algorithm was found to be racially biased, incorrectly "
            "flagging Black defendants as future criminals at twice the rate of white defendants.", body),
        Paragraph(
            "Privacy: Large-scale data collection required for AI training raises significant "
            "privacy concerns. The EU's General Data Protection Regulation (GDPR) introduced the "
            "'right to explanation' for automated decisions.", body),
        Paragraph(
            "Job Displacement: The World Economic Forum estimates AI and automation could displace "
            "85 million jobs by 2025, while creating 97 million new roles. The net impact depends "
            "heavily on retraining and education investments.", body),

        Paragraph("5. The Future of AI", h2),
        Paragraph(
            "Artificial General Intelligence (AGI) — AI that matches human cognitive ability across "
            "all domains — remains an open research problem. Current systems are narrow AI, excelling "
            "in specific tasks but lacking general reasoning. Leading AI labs including OpenAI, "
            "DeepMind, and Anthropic are actively researching paths toward AGI.", body),
        Paragraph(
            "Regulation is accelerating. The EU AI Act (2024) established the first comprehensive "
            "legal framework for AI, classifying systems by risk level and imposing transparency "
            "and accountability requirements.", body),
    ]

    doc.build(content)
    print(f"Created: {path}")


# ── Sample 2: Climate Change Report PDF ────────────────────────────────────────

def create_climate_report_pdf():
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except ImportError:
        return

    path = OUTPUT_DIR / "climate_change_report.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=72)

    styles = getSampleStyleSheet()
    h1, h2, body = styles["Heading1"], styles["Heading2"], styles["BodyText"]
    body.spaceAfter = 8

    content = [
        Paragraph("Climate Change: Scientific Consensus and Policy Responses", h1),
        Spacer(1, 0.2 * inch),

        Paragraph("Executive Summary", h2),
        Paragraph(
            "Global average temperatures have risen by approximately 1.1°C above pre-industrial "
            "levels. The Intergovernmental Panel on Climate Change (IPCC) Sixth Assessment Report "
            "(AR6) states with unequivocal certainty that human activities are the primary cause "
            "of this warming. Without immediate and drastic emissions reductions, warming is "
            "projected to exceed 1.5°C by the early 2030s.", body),

        Paragraph("1. Causes of Climate Change", h2),
        Paragraph(
            "The primary driver of contemporary climate change is the emission of greenhouse gases "
            "(GHGs) from human activities. Carbon dioxide (CO₂) from fossil fuel combustion accounts "
            "for approximately 76% of global GHG emissions. Methane (CH₄) — 25x more potent than "
            "CO₂ over 100 years — primarily comes from livestock, rice paddies, and natural gas leaks. "
            "Nitrous oxide (N₂O) from agriculture contributes a further 6%.", body),
        Paragraph(
            "Deforestation contributes about 10% of global emissions annually. Tropical forests "
            "absorb approximately 2.4 billion tonnes of carbon per year; their removal releases "
            "stored carbon while eliminating future absorption capacity.", body),

        Paragraph("2. Observed Impacts", h2),
        Paragraph(
            "Sea level rise has accelerated to 3.7mm per year (2006–2018), driven by thermal "
            "expansion and ice sheet melt. Greenland is losing 280 billion tonnes of ice annually; "
            "the West Antarctic Ice Sheet shows signs of irreversible destabilization.", body),
        Paragraph(
            "Extreme weather events have intensified. Heat waves are now 4.8x more likely than in "
            "the pre-industrial era. Atlantic hurricane intensity has increased, with Category 4 and "
            "5 storms becoming more frequent. Wildfire seasons have extended by 20% over the past "
            "three decades globally.", body),
        Paragraph(
            "Biodiversity loss is accelerating. Coral bleaching events that historically occurred "
            "once per 25–30 years now occur every 5–6 years. The Great Barrier Reef experienced "
            "mass bleaching in 2016, 2017, 2020, 2022, and 2024.", body),

        Paragraph("3. Mitigation Strategies", h2),
        Paragraph(
            "Energy Transition: Replacing fossil fuels with renewables is the single largest "
            "mitigation lever. Solar PV costs fell 89% between 2010 and 2020, making it the "
            "cheapest electricity source in history. Global renewable capacity must triple by 2030 "
            "to stay on a 1.5°C pathway.", body),
        Paragraph(
            "Carbon Capture and Storage (CCS): Direct air capture technologies can remove CO₂ "
            "directly from the atmosphere. Climeworks' Orca facility in Iceland captures 4,000 "
            "tonnes of CO₂ per year, though scaling to gigatonne levels requires cost reductions "
            "from $1,000/tonne today to under $100/tonne.", body),

        Paragraph("4. Paris Agreement and National Commitments", h2),
        Paragraph(
            "The Paris Agreement (2015) commits 195 countries to limiting warming to well below 2°C, "
            "pursuing 1.5°C. Countries submit Nationally Determined Contributions (NDCs) every five "
            "years. Current NDCs, if fully implemented, put the world on track for approximately "
            "2.5–2.9°C of warming by 2100 — far short of stated goals.", body),
        Paragraph(
            "The United States, EU, and China collectively account for 47% of global emissions. "
            "China, the world's largest emitter, has pledged carbon neutrality by 2060 and a peak "
            "in emissions before 2030. The US Inflation Reduction Act (2022) committed $369 billion "
            "to climate and clean energy investments.", body),
    ]

    doc.build(content)
    print(f"Created: {path}")


# ── Sample 3: Python Programming DOCX ─────────────────────────────────────────

def create_python_guide_docx():
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx not installed. Run: pip install python-docx")
        return

    path = OUTPUT_DIR / "python_programming_guide.docx"
    doc = Document()

    doc.add_heading("Python Programming: From Beginner to Intermediate", 0)

    doc.add_heading("Introduction", 1)
    doc.add_paragraph(
        "Python is a high-level, interpreted programming language created by Guido van Rossum "
        "and first released in 1991. Its design philosophy emphasizes code readability and "
        "simplicity. Python consistently ranks as the world's most popular programming language "
        "according to the TIOBE Index, GitHub statistics, and Stack Overflow developer surveys."
    )
    doc.add_paragraph(
        "Python's popularity stems from its versatility: it powers web backends (Django, FastAPI), "
        "data science (pandas, NumPy), machine learning (PyTorch, TensorFlow), automation scripts, "
        "and scientific computing. The language's 'batteries included' philosophy means its "
        "standard library covers most common tasks without external dependencies."
    )

    doc.add_heading("Core Data Types", 1)
    doc.add_paragraph(
        "Python has several built-in data types:\n"
        "- int: Integer numbers (e.g., 42, -17). Python integers have arbitrary precision.\n"
        "- float: Floating-point numbers (e.g., 3.14, -0.001). Uses IEEE 754 double precision.\n"
        "- str: Immutable Unicode strings. Support slicing, formatting, and rich methods.\n"
        "- list: Ordered, mutable sequences. Written as [1, 2, 3]. Support indexing and slicing.\n"
        "- tuple: Ordered, immutable sequences. Written as (1, 2, 3). Hashable.\n"
        "- dict: Key-value mappings. Written as {'key': 'value'}. O(1) average lookup.\n"
        "- set: Unordered collections of unique elements. Written as {1, 2, 3}."
    )

    doc.add_heading("Control Flow", 1)
    doc.add_paragraph(
        "Python uses indentation (typically 4 spaces) to define code blocks, replacing the "
        "braces used in C-family languages. This enforces readable code structure."
    )
    doc.add_paragraph(
        "The for loop iterates over any iterable: lists, strings, dictionaries, generators, "
        "and custom objects implementing __iter__. The range() function generates arithmetic "
        "progressions. List comprehensions provide a concise way to create lists: "
        "[x**2 for x in range(10) if x % 2 == 0]."
    )
    doc.add_paragraph(
        "Exception handling uses try/except/else/finally blocks. Python encourages EAFP "
        "(Easier to Ask Forgiveness than Permission) over LBYL (Look Before You Leap), "
        "preferring try/except over pre-condition checks."
    )

    doc.add_heading("Functions and Closures", 1)
    doc.add_paragraph(
        "Functions are first-class objects in Python: they can be assigned to variables, "
        "passed as arguments, and returned from other functions. This enables higher-order "
        "programming patterns like map(), filter(), and functools.reduce()."
    )
    doc.add_paragraph(
        "Python supports *args and **kwargs for variadic functions. Default argument values are "
        "evaluated once at function definition — a common gotcha when using mutable defaults "
        "like lists or dicts. Use None as a sentinel instead."
    )
    doc.add_paragraph(
        "Decorators (@decorator) are syntactic sugar for wrapping functions. They are widely "
        "used in frameworks: @app.route in Flask, @pytest.mark.parametrize in pytest, "
        "@property for computed attributes, and @functools.lru_cache for memoization."
    )

    doc.add_heading("Object-Oriented Programming", 1)
    doc.add_paragraph(
        "Python supports multiple inheritance, duck typing, and operator overloading. "
        "Special methods (__dunder__ methods) allow classes to integrate with Python's "
        "protocols: __len__, __iter__, __getitem__ for container behaviour; "
        "__enter__ and __exit__ for context managers; __add__, __mul__ for arithmetic."
    )
    doc.add_paragraph(
        "Dataclasses (@dataclass decorator, Python 3.7+) reduce boilerplate for classes "
        "that primarily store data, auto-generating __init__, __repr__, and __eq__. "
        "Pydantic extends this with runtime type validation, making it popular for API schemas."
    )

    doc.add_heading("Python Ecosystem and Package Management", 1)
    doc.add_paragraph(
        "pip is Python's standard package installer. Virtual environments (venv or virtualenv) "
        "isolate project dependencies. pyproject.toml (PEP 518) has replaced setup.py as the "
        "standard build specification format. Poetry and uv are modern dependency managers "
        "that handle lock files and virtual environment management automatically."
    )
    doc.add_paragraph(
        "PyPI (Python Package Index) hosts over 500,000 packages. Key ecosystem libraries: "
        "requests/httpx for HTTP, pandas for tabular data, NumPy for numerical computing, "
        "matplotlib/seaborn for visualization, SQLAlchemy for databases, and FastAPI for APIs."
    )

    doc.add_heading("Performance Considerations", 1)
    doc.add_paragraph(
        "CPython (the reference implementation) uses a Global Interpreter Lock (GIL) that "
        "prevents true thread-level parallelism for CPU-bound tasks. For CPU parallelism, "
        "use the multiprocessing module or process pools. For I/O-bound concurrency, "
        "asyncio with async/await provides cooperative multitasking with excellent performance."
    )
    doc.add_paragraph(
        "Python 3.13 introduces an experimental free-threaded mode (--disable-gil) that "
        "enables true multi-core parallelism. This is expected to become the default in "
        "Python 3.15. NumPy and other C-extension libraries release the GIL during "
        "computation, achieving near-native performance today."
    )

    doc.save(str(path))
    print(f"Created: {path}")


if __name__ == "__main__":
    print("Generating sample documents...")
    create_ai_overview_pdf()
    create_climate_report_pdf()
    create_python_guide_docx()
    print("\nDone. Files written to:", OUTPUT_DIR)
